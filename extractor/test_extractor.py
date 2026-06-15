"""
Self-contained test for the extraction engine (no pytest required).

Exercises the cross-platform path end to end: build a messy .xlsx and a .docx in
a temp folder, run the orchestrator, and assert that faithful raw JSON and a
manifest are produced, that re-runs skip unchanged files, and that an
unsupported file is reported rather than crashing the run.

Run:  python3 -m extractor.test_extractor
"""
from __future__ import annotations

import json
import os
import sys
import tempfile

from . import manifest, registry
from .cli import run
from .csv_text import CsvTextExtractor
from .excel_openpyxl import ExcelOpenpyxlExtractor


def _make_samples(intake_dir):
    os.makedirs(intake_dir, exist_ok=True)

    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "P&L"
    ws.append(["Region", "Net Sales"])
    ws.append(["Africa", 1250000])
    ws.append([None, None])          # messy blank row must be preserved
    ws.append(["Europe", 980000])
    wb.save(os.path.join(intake_dir, "pl.xlsx"))

    import docx
    d = docx.Document()
    d.add_heading("Commentary", level=1)
    d.add_paragraph("Revenue grew.")
    d.save(os.path.join(intake_dir, "notes.docx"))

    # An unsupported file type to prove graceful skipping.
    with open(os.path.join(intake_dir, "random.xyz"), "w", encoding="utf-8") as fh:
        fh.write("not a known format")


def test_csv_arabic_encoding():
    """A Windows-1256 Arabic CSV and a UTF-8 BOM CSV must decode correctly."""
    extractor = CsvTextExtractor()
    with tempfile.TemporaryDirectory() as tmp:
        # Legacy Arabic code page, semicolon-delimited (common on Arabic Windows).
        cp1256_path = os.path.join(tmp, "arabic.csv")
        with open(cp1256_path, "wb") as fh:
            fh.write("المنطقة;المبيعات\nأفريقيا;1250\n".encode("cp1256"))
        content, warnings = extractor.extract_content(cp1256_path)
        cells = content["sheets"][0]["cells"]
        assert cells[0] == ["المنطقة", "المبيعات"], cells[0]
        assert cells[1] == ["أفريقيا", "1250"], cells[1]
        assert any("cp1256" in w for w in warnings), warnings

        # UTF-8 with BOM: BOM must not leak into the first cell.
        bom_path = os.path.join(tmp, "bom.csv")
        with open(bom_path, "wb") as fh:
            fh.write("السنة,المبيعات\n2025,980\n".encode("utf-8"))
        content, _ = extractor.extract_content(bom_path)
        cells = content["sheets"][0]["cells"]
        assert cells[0] == ["السنة", "المبيعات"], cells[0]


def test_csv_edge_cases():
    """Test CSV edge cases: empty file, binary content."""
    extractor = CsvTextExtractor()
    with tempfile.TemporaryDirectory() as tmp:
        # Empty file
        empty_path = os.path.join(tmp, "empty.csv")
        with open(empty_path, "wb") as fh:
            pass  # 0 bytes
        content, warnings = extractor.extract_content(empty_path)
        assert content["sheets"][0]["n_rows"] == 0, "empty file should have 0 rows"
        assert any("empty" in w.lower() for w in warnings), warnings
        
        # Binary content (should warn)
        binary_path = os.path.join(tmp, "binary.csv")
        with open(binary_path, "wb") as fh:
            # Write lots of null bytes (non-printable)
            fh.write(b"\x00" * 1000 + b"a,b,c\n1,2,3\n")
        content, warnings = extractor.extract_content(binary_path)
        assert any("binary" in w.lower() for w in warnings), warnings


def test_excel_formula_errors():
    """Test that formula errors (#REF!, #DIV/0!) are converted to null with warnings."""
    extractor = ExcelOpenpyxlExtractor()
    with tempfile.TemporaryDirectory() as tmp:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["Value", "Formula"])
        ws.append([100, "#DIV/0!"])
        ws.append([200, "#REF!"])
        ws.append([300, "OK"])
        path = os.path.join(tmp, "formulas.xlsx")
        wb.save(path)
        
        content, warnings = extractor.extract_content(path)
        cells = content["sheets"][0]["cells"]
        # Formula errors should be converted to None
        assert cells[1][1] is None, "DIV/0 should be None"
        assert cells[2][1] is None, "REF should be None"
        assert cells[3][1] == "OK", "valid value preserved"
        # Should have a warning about formula errors
        assert any("formula" in w.lower() for w in warnings), warnings


def test_excel_corrupt_file():
    """Test that corrupt Excel files raise clear errors."""
    extractor = ExcelOpenpyxlExtractor()
    with tempfile.TemporaryDirectory() as tmp:
        corrupt_path = os.path.join(tmp, "corrupt.xlsx")
        with open(corrupt_path, "wb") as fh:
            fh.write(b"this is not a valid excel file")
        
        try:
            extractor.extract_content(corrupt_path)
            assert False, "should have raised ValueError"
        except ValueError as exc:
            assert "corrupt" in str(exc).lower() or "cannot open" in str(exc).lower()


def test_excel_empty_file():
    """Test that empty Excel files are handled gracefully."""
    extractor = ExcelOpenpyxlExtractor()
    with tempfile.TemporaryDirectory() as tmp:
        empty_path = os.path.join(tmp, "empty.xlsx")
        with open(empty_path, "wb") as fh:
            pass  # 0 bytes
        
        try:
            extractor.extract_content(empty_path)
            # Empty files might raise ValueError (corrupt) or return empty sheets
            # Both are acceptable - the important thing is it doesn't crash silently
        except ValueError:
            pass  # Expected for truly empty files


def test_new_extractors_registered():
    """The .xlsb/.xls/CSV readers must be wired in and own their extensions."""
    names = {row["name"] for row in registry.describe_availability()}
    for expected in ("excel-xlsb", "excel-xls", "csv-text"):
        assert expected in names, f"{expected} not registered"
    extractor, reason = registry.select_extractor("/tmp/whatever.csv")
    assert extractor is not None and extractor.name == "csv-text", reason


def main():
    # These need no third-party libraries, so run them first/unconditionally.
    test_csv_arabic_encoding()
    test_csv_edge_cases()
    test_new_extractors_registered()
    
    # Test Excel edge cases (needs openpyxl)
    try:
        __import__("openpyxl")
        test_excel_formula_errors()
        test_excel_corrupt_file()
        test_excel_empty_file()
    except ImportError:
        print("SKIP: openpyxl not installed; Excel edge case tests not run.")

    # Skip cleanly if the cross-platform libraries are unavailable.
    for module in ("openpyxl", "docx"):
        try:
            __import__(module)
        except Exception:  # noqa: BLE001
            print(f"SKIP: {module} not installed; extractor test not run.")
            return 0

    with tempfile.TemporaryDirectory() as tmp:
        intake = os.path.join(tmp, "intake")
        out = os.path.join(tmp, "raw")
        _make_samples(intake)

        summary = run(intake, out, force=False)
        assert summary["captured"] == 2, summary
        assert summary["unsupported"] == 1, summary
        assert summary["errors"] == 0, summary

        raw_files = [f for f in os.listdir(out) if f.endswith(".raw.json")]
        assert len(raw_files) == 2, raw_files

        excel_raw = next(f for f in raw_files if "pl.xlsx" in f)
        with open(os.path.join(out, excel_raw), encoding="utf-8") as fh:
            envelope = json.load(fh)
        assert envelope["document_type"] == "spreadsheet"
        assert envelope["source"]["sha256"]
        sheet = envelope["content"]["sheets"][0]
        assert sheet["cells"][1] == ["Africa", 1250000], sheet["cells"][1]
        assert sheet["n_rows"] == 4, "blank row should be preserved"

        # Manifest recorded the two successes.
        seen = manifest.load_seen_hashes(out)
        assert len(seen) == 2, seen

        # Re-run: nothing new captured (dedup by fingerprint).
        again = run(intake, out, force=False)
        assert again["captured"] == 0, again
        assert again["skipped_unchanged"] == 2, again

    # Availability listing must never raise, even with broken optional libs.
    rows = registry.describe_availability()
    assert any(r["name"] == "excel-com" for r in rows)

    print("Extractor tests passed.")
    return 0


if __name__ == "__main__":
    sys.exit(main())

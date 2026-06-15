"""
word_docx.py - Word extractor (cross-platform, testable).

Reads .docx files with python-docx: paragraphs (with their style) and tables.
This runs anywhere. On Windows a COM-based Word extractor could be added for
legacy .doc files and richer fidelity, mirroring excel_com.py; the envelope
shape would stay identical.
"""
from __future__ import annotations

import os

from .base import Extractor, optional_import


class WordDocxExtractor(Extractor):
    name = "word-docx"
    document_type = "document"
    extensions = (".docx",)

    def is_available(self):
        if optional_import("docx") is None:
            return False, "python-docx is not installed (pip install python-docx)"
        return True, "ok"

    def extract_content(self, path: str):
        import docx

        warnings = []
        
        # Guard: file size limit (50MB)
        file_size = os.path.getsize(path)
        if file_size > 50 * 1024 * 1024:
            warnings.append(
                f"File is very large ({file_size / (1024*1024):.1f} MB). "
                "Extraction may be slow or incomplete."
            )
        
        # Guard: empty file
        if file_size == 0:
            warnings.append("File is empty (0 bytes).")
            return {"paragraphs": [], "tables": []}, warnings
        
        try:
            document = docx.Document(path)
        except Exception as exc:
            raise ValueError(
                f"Cannot open document: {exc}. File may be corrupt, "
                "encrypted, or not a valid .docx file."
            ) from exc

        paragraphs = []
        parse_errors = 0
        for idx, p in enumerate(document.paragraphs, start=1):
            try:
                if p.text and p.text.strip():
                    paragraphs.append({
                        "style": p.style.name if p.style else None,
                        "text": p.text
                    })
            except Exception as exc:
                parse_errors += 1
                if parse_errors <= 3:
                    warnings.append(f"Paragraph {idx}: could not read ({exc})")
        
        if parse_errors > 3:
            warnings.append(f"... and {parse_errors - 3} more paragraph errors")
        
        tables = []
        for table_idx, table in enumerate(document.tables, start=1):
            try:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                tables.append({"rows": rows})
            except Exception as exc:
                warnings.append(f"Table {table_idx}: could not read ({exc})")

        if not paragraphs and not tables:
            warnings.append("Document had no readable text or tables.")
        return {"paragraphs": paragraphs, "tables": tables}, warnings
